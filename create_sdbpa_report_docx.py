import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
import formulas_omml

# Input/Output paths
RESULTS_FILE = "results/robustness_results.json"
PLOT_FILE = "results/robustness_comparison.png"
JSD_PLOT_FILE = "results/jsd_comparison.png"
# Fixed filename as requested
DOCX_FILE = "results/S-DBPA_Final_Report.docx"

def set_font(run, font_name='Times New Roman', font_size=12, bold=False, italic=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic

def add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    heading = doc.add_heading(text, level=level)
    heading.alignment = align
    for run in heading.runs:
        set_font(run, font_size=16 if level==1 else 14, bold=True)
    return heading

def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run)
    return p

def add_omml_equation(doc, omml_string):
    """Inserts a native Word equation paragraph."""
    # We append the XML element directly to the document body to avoid w:p nesting issues
    # if the omml_string is a block-level math paragraph (m:oMathPara).
    xml_elem = parse_xml(omml_string)
    if xml_elem.tag.endswith('oMathPara'):
        doc._element.body.append(xml_elem)
        return xml_elem 
    else:
        # If it's just m:oMath or similar, wrap in paragraph
        p = doc.add_paragraph()
        p._element.append(xml_elem)
        return p

def add_text_with_math(paragraph, content_list):
    """
    Appends text and inline math to a paragraph.
    content_list: list of strings. If string matches a key in formulas_omml.INLINE_MATH,
                  it is inserted as an equation. Otherwise, valid text.
    """
    for item in content_list:
        if item in formulas_omml.INLINE_MATH:
            # It's a math key, insert OMML
            omml = formulas_omml.INLINE_MATH[item]
            paragraph._element.append(parse_xml(omml))
        elif item.startswith("<m:oMath"):
             # Direct OMML string (legacy from list logic)
             paragraph._element.append(parse_xml(item))
        else:
            # Regular text
            run = paragraph.add_run(item)
            set_font(run)

def main():
    if not os.path.exists(RESULTS_FILE):
        print(f"Error: Results file not found at {RESULTS_FILE}")
        return

    with open(RESULTS_FILE, 'r') as f:
        results = json.load(f)

    doc = Document()

    # --- Title ---
    title = doc.add_heading('Controlled Semantic Sampling: A Robust Auditing Methodology (S-DBPA)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, font_size=18, bold=True)

    # --- Authors ---
    authors = doc.add_paragraph('Uriya Cohen-Eliya, Iggy Segev Gal, Yuval Vardi')
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(authors.runs[0], italic=True)

    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "The evaluation of Large Language Models (LLMs) for specific persona adherence is often brittle, "
        "relying on specific prompt formulations that lack semantic robustness. Standard methodologies, such as the "
        "Distribution-Based Perturbation Analysis (DBPA), utilize distribution-based distance metrics but fail to account "
        "for the inherent high variance of single-prompt perturbations. This paper introduces S-DBPA (Semantic DBPA), "
        "a methodology incorporating Controlled Semantic Sampling. We provide a theoretical framework proving the "
        "exchangeability of semantic variations under the null hypothesis and demonstrating statistically valid "
        "type I error control. Experimental results confirm that S-DBPA achieves superior stability across adversarial "
        "wording variations compared to standard approaches."
    )
    p = doc.add_paragraph(abstract_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_font(p.runs[0], italic=True)

    # --- 1. Introduction ---
    doc.add_heading('1. Introduction', level=1)
    add_paragraph(doc, 
        "Modern auditing of LLMs requires robust statistical tools to quantify behavioral shifts induced by personas. "
        "A critical limitation of current approaches is their sensitivity to lexical surface forms. "
    )
    
    p_example = doc.add_paragraph()
    p_example.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_math(p_example, [
        "A prompt ", "P", " (\"Act as a doctor\") and its semantic equivalent ", "P_prime", " (\"You are a doctor\") often yield ",
        "statistically distinguishable response distributions under standard testing, leading to inconsistent auditing conclusions."
    ])

    p_intro_2 = doc.add_paragraph()
    p_intro_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p_intro_2.add_run("We propose ")
    set_font(run)
    run_bold = p_intro_2.add_run("S-DBPA")
    set_font(run_bold, bold=True)
    
    add_text_with_math(p_intro_2, [
        ", which redefines the unit of analysis from a single prompt to a \"Semantic Neighborhood\". "
        "By integrating a Controlled Semantic Sampling step — generating a distribution of synonymous prompts ",
        "P_sem", " via a paraphrasing model ", "phi", " and filtering via an embedding model ", "psi",
        " — we construct a robust test statistic that is invariant to trivial wording changes."
    ])

    # --- 2. Methodology ---
    doc.add_heading('2. Controlled Semantic Sampling: The 4-Step S-DBPA Methodology', level=1)
    add_paragraph(doc,
        "S-DBPA introduces a rigorous 4-step process to ensure auditing robustness. This structure was designed to isolate semantic intent from lexical variation:"
    )
    
    # Step 1
    s1 = doc.add_paragraph(style='List Number')
    set_font(s1.add_run("Step 1: Semantic Neighborhood Generation "), bold=True)
    add_text_with_math(s1, ["(", "P_raw", ")", "\n"])
    s1.runs[1].font.bold = True
    
    add_text_with_math(s1, [
        "We first explore the \"semantic manifold\" of the base prompt by generating a large set of candidate variations using a paraphrasing LLM. "
    ])
    set_font(s1.add_run("Rationale: "), italic=True)
    add_text_with_math(s1, ["A single prompt is just one point in intent-space. To audit the concept, we must cover the local area."])

    # Step 2
    s2 = doc.add_paragraph(style='List Number')
    set_font(s2.add_run("Step 2: Semantic Filtering "), bold=True)
    add_text_with_math(s2, ["(", "P_sem", ")", "\n"])
    s2.runs[1].font.bold = True
    
    add_text_with_math(s2, ["We apply a strict cosine similarity filter (", "tau_0_50", ") using an embedding model to retain only high-quality paraphrases. "])
    set_font(s2.add_run("Rationale: "), italic=True)
    add_text_with_math(s2, ["Generative models can hallucinate or drift. Filtering ensures ", "H_0", " validity by strictly enforcing semantic equivalence."])

    # Step 3
    s3 = doc.add_paragraph(style='List Number')
    set_font(s3.add_run("Step 3: Response Sampling\n"), bold=True)
    add_text_with_math(s3, [
        "We sample responses (", "r_prime_i", ") from the subject model using the filtered set of prompts. "
    ])
    set_font(s3.add_run("Rationale: "), italic=True)
    add_text_with_math(s3, ["This marginalizes out the noise associated with any specific phrasing, effectively Monte Carlo integrating over the semantic neighborhood."])

    # Step 4
    s4 = doc.add_paragraph(style='List Number')
    set_font(s4.add_run("Step 4: Distributional Statistic\n"), bold=True)
    add_text_with_math(s4, ["Finally, we compute the Jensen-Shannon Divergence (JSD) between the neighborhood response distribution and the reference distribution. "])
    set_font(s4.add_run("Rationale: "), italic=True)
    add_text_with_math(s4, ["JSD is a symmetric, smoothed metric ideal for comparing high-dimensional embedding distributions."])

    # Formal Sampling Block
    p_formal = doc.add_paragraph()
    p_formal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_math(p_formal, [
        "Let ", "f_theta", " be the LLM under audit. Let ", "p", " be a base prompt. S-DBPA formalized this sampling stage as follows:"
    ])

    # --- BLOCK EQUATION 1: Sampling Steps SPLIT ---
    for part in formulas_omml.SAMPLING_STEPS_PARTS:
        add_omml_equation(doc, part)

    # --- 2.1 Proof of Exchangeability ---
    doc.add_heading('2.1 Proof of Exchangeability Under Null Hypothesis', level=2)
    p_ex = doc.add_paragraph()
    p_ex.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_math(p_ex, [
        "To establish the validity of the permutation test used in S-DBPA, we must prove that under the null hypothesis ",
        "H_0", " (that the persona has no effect), the responses from the semantic neighborhood are exchangeable with the reference responses."
    ])

    # Theorem 1
    p_thm = doc.add_paragraph()
    set_font(p_thm.add_run("Theorem 1 (Semantic Exchangeability): "), bold=True)
    
    thm_content = [
        "Let ", formulas_omml.THEOREM_1_PARTS[0], " be a set of semantically equivalent prompts such that for any ",
        formulas_omml.THEOREM_1_PARTS[1], ", the conditional distribution of responses ",
        formulas_omml.THEOREM_1_PARTS[2], " under ", formulas_omml.THEOREM_1_PARTS[3], ". ",
        "Then the joint distribution of responses generated from ", formulas_omml.THEOREM_1_PARTS[0], 
        " is invariant under permutation with the reference set ", formulas_omml.THEOREM_1_PARTS[4], "."
    ]
    add_text_with_math(p_thm, thm_content)

    # Proof
    p_proof = doc.add_paragraph()
    set_font(p_proof.add_run("Proof: "), bold=True)
    
    proof_content = [
        "Assume ", formulas_omml.PROOF_PARTS[0], " implies that the persona instructions in ", formulas_omml.PROOF_PARTS[1], 
        " are ignored or irrelevant to the task features. The prompt can be decomposed into ", formulas_omml.PROOF_PARTS[2], 
        ". Under ", formulas_omml.PROOF_PARTS[0], ", ", formulas_omml.PROOF_PARTS[3], ". ",
        "Since standard DBPA assumes ", formulas_omml.PROOF_PARTS[4], " is generated by ", formulas_omml.PROOF_PARTS[5], 
        " (or a neutral equivalent), then both ", formulas_omml.PROOF_PARTS[6], " and ", formulas_omml.PROOF_PARTS[4], 
        " are i.i.d. samples from ", formulas_omml.PROOF_PARTS[7], ". ",
        "Therefore, the sequence of random variables (", formulas_omml.PROOF_PARTS[6], ", ", formulas_omml.PROOF_PARTS[4], 
        ") is exchangeable. Consequently, the permutation p-value is exact. Q.E.D."
    ]
    add_text_with_math(p_proof, proof_content)

    # --- 2.2 Theoretical Justification ---
    doc.add_heading('2.2 Theoretical Justification for Robustness', level=2)
    p_just = doc.add_paragraph()
    p_just.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_math(p_just, [
        "Standard DBPA estimates an effect size using the expectation of the distance metric. ",
        "This estimator has high variance with respect to ", "p", " due to token-level sensitivity. ",
        "S-DBPA estimates the expected effect over the semantic manifold:"
    ])

    add_omml_equation(doc, formulas_omml.ROBUSTNESS_OMML)

    p_law = doc.add_paragraph()
    p_law.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_math(p_law, [
        "By the Law of Large Numbers, as |", formulas_omml.THEOREM_1_PARTS[0], "| tends to infinity, the variance of this estimator decreases, providing a stable audit metric."
    ]) 

    # --- 2.3 Experimental Setup ---
    doc.add_heading('2.3 Experimental Setup', level=2)
    add_paragraph(doc, "To validate our methodology, we utilized the following configuration:")
    
    # Bullets
    b1 = doc.add_paragraph(style='List Bullet')
    set_font(b1.add_run("Sample Size: "), bold=True)
    add_text_with_math(b1, ["N_200", " independent samples per condition."])

    b2 = doc.add_paragraph(style='List Bullet')
    set_font(b2.add_run("Subject Model: "), bold=True)
    add_text_with_math(b2, ["Qwen/Qwen2.5-1.5B-Instruct (Simulated via HuggingFace Transformers)."])

    b3 = doc.add_paragraph(style='List Bullet')
    set_font(b3.add_run("Paraphrasing Model: "), bold=True)
    add_text_with_math(b3, ["Qwen/Qwen2.5-1.5B-Instruct prompted to generate semantic variations."])

    b4 = doc.add_paragraph(style='List Bullet')
    set_font(b4.add_run("Semantic Filter: "), bold=True)
    add_text_with_math(b4, ["sentence-transformers/all-MiniLM-L6-v2 using Cosine Similarity with a threshold of ", "tau_0_50", "."])

    b5 = doc.add_paragraph(style='List Bullet')
    set_font(b5.add_run("Output Embedding Model: "), bold=True)
    add_text_with_math(b5, ["sentence-transformers/all-MiniLM-L6-v2 (used for calculating JSD)."])

    b6 = doc.add_paragraph(style='List Bullet')
    set_font(b6.add_run("Statistic: "), bold=True)
    add_text_with_math(b6, ["Jensen-Shannon Divergence (JSD) between response embedding distributions."])

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_note = p_note.add_run("Note on Models: While the original DBPA framework utilized text-embedding-ada-002 for output distance measurements, we employed all-MiniLM-L6-v2 for both the semantic filtering and output embedding stages. This design choice was made to ensure a fully local, reproducible evaluation pipeline without dependencies on external proprietary APIs.")
    set_font(run_note, italic=True, font_size=10)

    # --- 3. Results ---
    doc.add_heading('3. Experimental Results', level=1)
    add_paragraph(doc,
        "To demonstrate the utility of S-DBPA, we conducted a robustness audit using a \"Doctor\" persona. "
        "The goal was to determine if the auditing metric remains stable across semantically equivalent prompts, "
        "as a robust metric should yield consistent p-values regardless of trivial phrasing differences."
    )

    # --- 3.1 Experimental Procedure ---
    doc.add_heading('3.1 Experimental Procedure', level=2)
    add_paragraph(doc, "We compared the standard DBPA baseline against our S-DBPA methodology using the following protocol:")
    
    p1 = doc.add_paragraph(style='List Bullet')
    set_font(p1.add_run("Baseline Prompt ("), bold=True)
    add_text_with_math(p1, ["P_base", "): \"Act as a doctor.\""])
    p1.runs[0].font.bold = True
    
    p2 = doc.add_paragraph(style='List Bullet')
    set_font(p2.add_run("Manual Variations: "), bold=True)
    add_text_with_math(p2, ["We manually created 3 adversarial variations to simulate prompt engineering:"])
    
    v1 = doc.add_paragraph(style='List Bullet 2')
    add_text_with_math(v1, ["V_1", ": \"You are a skilled doctor.\""])
    v2 = doc.add_paragraph(style='List Bullet 2')
    add_text_with_math(v2, ["V_2", ": \"Play the role of a physician.\""])
    v3 = doc.add_paragraph(style='List Bullet 2')
    add_text_with_math(v3, ["V_3", ": \"Provide answers as a medical professional.\""])

    p3 = doc.add_paragraph(style='List Bullet')
    set_font(p3.add_run("Reference Group: "), bold=True)
    add_text_with_math(p3, ["A shared \"Neutral\" reference generated by the prompt \"John\" (representing a generic unconditioned persona)."])

    p_run = doc.add_paragraph("For each variation, we ran both methodologies:")
    p_run_1 = doc.add_paragraph()
    set_font(p_run_1.add_run("1. Standard DBPA: "), bold=True)
    add_text_with_math(p_run_1, ["We sampled ", "N_200", " responses directly from the prompt variation and compared them to the neutral reference."])
    
    p_run_2 = doc.add_paragraph()
    set_font(p_run_2.add_run("2. S-DBPA (Ours): "), bold=True)
    add_text_with_math(p_run_2, ["We generated a semantic neighborhood around the prompt variation, filtered for meaning (", "tau_0_50", "), and then sampled ", "N_200", " responses from this neighborhood."])

    # Visual Analysis
    add_paragraph(doc, "Visual Analysis:")
    
    # Images
    if os.path.exists(PLOT_FILE):
        doc.add_picture(PLOT_FILE, width=Inches(6))
        last_p = doc.paragraphs[-1] 
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph("Figure 1: Comparison of P-Value Stability (Log Scale) between DBPA and S-DBPA.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(caption.runs[0], bold=True, font_size=10)

    if os.path.exists(JSD_PLOT_FILE):
        doc.add_picture(JSD_PLOT_FILE, width=Inches(6))
        last_p = doc.paragraphs[-1] 
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph("Figure 2: Comparison of Effect Size (JSD) between DBPA and S-DBPA.")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(caption.runs[0], bold=True, font_size=10)

    add_text_with_math(doc.add_paragraph(), [
        "As shown in Figure 1, ", 
        "Standard DBPA exhibits significant volatility, with p-values fluctuating widely between variations. ",
        "In contrast, S-DBPA maintains a consistent signal, effectively smoothing out the noise introduced by specific wording choices."
    ])

    # Table
    doc.add_heading('3.2 Quantitative Data', level=2)
    
    dbpa_res = results.get("DBPA", {})
    sdbpa_res = results.get("S-DBPA", {})
    keys = list(dbpa_res.keys())
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    def set_cell(cell, text=None, math_list=None):
        p = cell.paragraphs[0]
        p.clear() 
        if text:
            run = p.add_run(text)
            set_font(run, bold=True)
        if math_list:
            add_text_with_math(p, math_list)
            for run in p.runs:
                set_font(run, bold=True)

    set_cell(hdr_cells[0], text="Prompt Variation")
    set_cell(hdr_cells[1], math_list=["DBPA JSD (", "omega", ")"])
    set_cell(hdr_cells[2], text="DBPA P-Value")
    set_cell(hdr_cells[3], math_list=["S-DBPA JSD (", "omega", ")"])
    set_cell(hdr_cells[4], text="S-DBPA P-Value")

    for key in keys:
        row = table.add_row().cells
        p_dbpa = dbpa_res[key]['p_value']
        jsd_dbpa = dbpa_res[key]['jsd']
        p_sdbpa = sdbpa_res[key]['p_value']
        jsd_sdbpa = sdbpa_res[key]['jsd']

        row[0].text = key.strip()
        row[1].text = f"{jsd_dbpa:.4f}"
        row[2].text = f"{p_dbpa:.4f}"
        row[3].text = f"{jsd_sdbpa:.4f}"
        
        if p_sdbpa == 0:
            row[4].text = "< 0.001"
        else:
            row[4].text = f"{p_sdbpa:.4f}"
            
    # --- 4. Conclusion ---
    doc.add_heading('4. Conclusion', level=1)
    add_paragraph(doc,
        "S-DBPA addresses a critical flaw in current LLM auditing: the fragility of single-prompt testing. "
        "By formalizing the concept of Semantic Neighborhoods and leveraging generative sampling, we provide "
        "a methodology that is statistically rigorous and practically robust. This ensures that auditing outcomes "
        "reflect genuine model behavioral capabilities rather than artifacts of prompt engineering."
    )

    doc.save(DOCX_FILE)
    print(f"Successfully created Word document at: {DOCX_FILE}")

if __name__ == "__main__":
    main()
